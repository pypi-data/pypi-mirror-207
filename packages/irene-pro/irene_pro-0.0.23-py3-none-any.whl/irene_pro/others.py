
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from time import strftime
from datetime import datetime
import pyperclip as copier
import os, shutil
import qrcode

import openpyxl
import win32print
import win32api
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import win32com.client as win32
import shutil
import pandas as pd
from socket import gaierror
from threading import Thread

class make_thread(Thread):
    def __init__(self, target, **kwargs):
        super().__init__(target=target, **kwargs)
        self.start()

class DateTime:
    def __init__(self, week = datetime.today().isocalendar()[1], day = strftime("%Y-%m-%d"), 
                 month = strftime("%Y-%m"), year = strftime("%Y"), current_weeks = None, combined = int(strftime("%Y%m%d%H%M")),
                   combined_date = None) -> None:
        self.week = week
        self.day = day
        self.month = month
        self.year = year
        self.current_weeks = current_weeks
        self.current_weeks = str(week) + str(year)
        self.combined = combined
        self.combined_date = combined_date

    def datedelta(self, start, end):
        date1 = datetime.strptime(start, "%Y-%m-%d %H:%M:%S")
        date2 = datetime.strptime(end, "%Y-%m-%d %H:%M:%S")
        delta = date2 - date1
        return delta
    @property
    def combinedDate(self):
        return [self.year, self.month, self.current_weeks, self.day, self.combined]

class folders:
    def __init__(self, folder_path = None) -> None:
        self.dir = folder_path

    def make_dir(self):
        try:
            os.makedirs(self.dir)
        except Exception:
            pass
        finally:
            return self.dir
    
    def cut_file(self, displayer, combined_time:str, source_file_path = None, destiny_path = None, destiny_filename = None):
        try:
            if source_file_path and destiny_path:
                if not destiny_filename:
                    destiny_filename = os.path.basename(source_file_path)
                # Cut the file and paste it in the destination directory
                extension = os.path.splitext(source_file_path)[1]
                shutil.move(source_file_path, os.path.join(destiny_path, destiny_filename + f"@{combined_time}@" + extension))
        except Exception:
            displayer.showerror("Failed", "Something went wrong")

class Create:
    def __init__(self, text, filename):
        self.text = text
        self.filename = filename

    def qrcode(self):
        qr = qrcode.make(self.text)
        qr.save(self.filename)

class doc:
    def __init__(self, machine, logo = None,header_title = "Issue Report", heading_title = f"Issue occurence report {strftime('%Y-%m-%d')}", footer = f"ProcessPro © {strftime('%Y')}", filename = "IssueApt2023.docx", qrcode  =None) -> None:
        self.logo = logo
        self.qrcode = qrcode
        self.heading_title = heading_title
        self.header_title = header_title
        self.footer = footer
        self.filename = filename
        self.machine = machine
        self.emoj = "\u26a0\ufe0f"*30
        self.heading = f"{self.emoj[:int(len(self.emoj)/2)]}Issue at {self.machine}{self.emoj[int(len(self.emoj)/2):]}"
        self.document = Document()
        # self.document.add_picture("update.png", width = Inches(.1))
        
    def Heading(self, levels = {'zone':'zone1', 'section':'section1', 'line':'line2', 'machine':'machine1', 'machine part':'screw', "writer":"Irene", "Issue date":"2023-06-20", "Record date":"2023-06-24"}):
        """levels = {'zone':'zone1', 'section':'section1', 'line':'line2', 'machine':'machine1', 'machine part':'screw'}"""
        # HEADING 
        print("LEVELS: ", levels)
        header = self.document.sections[0].header
        htable=header.add_table(1, 2, Inches(6))
        htab_cells=htable.rows[0].cells
        ht0=htab_cells[0].add_paragraph()
        kh=ht0.add_run()
        # IMAGES IN HEADER
        if self.logo:
            kh.add_picture(self.logo, width=Inches(1.55), height = Inches(1.2))
        if self.qrcode:
            kh.add_picture(self.qrcode, width=Inches(1))
        # HEADER TITLE
        ht1=htab_cells[1].add_paragraph(self.header_title)
        # HEADING OF THE DOCUMENT
        table_title = list(levels.keys())
        table_value = list(levels.values())
        
        table = self.document.add_table(rows = 2, cols = len(levels))
        i = 0
        for tit in table_title:
            c = table.cell(0, i)
            c.text = tit
            i += 1
        j = 0
        for content in table_value:
            b = table.cell(1, j)
            b.text = content
            j += 1
        table.style = "Table Grid"
            
        self.document.add_paragraph("\n")

        # align heading of the document
        ht1.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # FOOTER 
        if self.footer:
            f = self.document.sections[0].footer
            f.paragraphs[0].text = self.footer
    
    def body(self, description, image = None):
        self.document.add_paragraph(description)
        p = self.document.add_paragraph()
        r = p.add_run()
        if image and os.path.isfile(image):
            r.add_picture(image, width = Inches(6))
            
    def team(self, team_members = [], title = "Team to handle the issue"):
        if title and team_members:
            team_title = self.document.add_heading(title, 2)
            team_title.alignment = WD_ALIGN_PARAGRAPH.LEFT

            team_table = self.document.add_table(cols = len(team_members[0]), rows = len(team_members))
            row = 0
            for member in team_members:
                for col in range(len(member)):
                    Cell = team_table.cell(row, col)
                    Cell.text = member[col]
                row += 1
            team_table.style = "Table Grid"

    def save_doc(self):
        self.document.save(self.filename)

class Email:
    def __init__(self):
        pass

    def classify_email_address(self, all_selected_email_addresses):
        gmail = []
        outlook = []
        for mail in all_selected_email_addresses:
            if "gmail" in mail:
                gmail.append(mail)
            else:
                outlook.append(mail)
        return gmail, outlook
    
    def sendGmail(self, receivers = [], sender = "irene.study.2023@gmail.com", subject = "Document for the last report", body = "Report", filenames = [], password = "qzpcgergzvjydryc"):
        len_receivers = len(receivers)
        sent = 0
        fromaddr = sender
        # toaddr = "ing.rw.evode@gmail.com"
        
        toaddr = ", ".join(receivers)
        # instance of MIMEMultipart
        msg = MIMEMultipart()

        # storing the senders email address
        msg['From'] = fromaddr

        # storing the receivers email address
        msg['To'] = ", ".join(receivers)

        # storing the subject
        msg['Subject'] = subject

        # string to store the body of the mail
        body = body

        # attach the body with the msg instance
        msg.attach(MIMEText(body, 'plain'))

        if filenames:
            for filename in filenames:
                attachment = open(filename, "rb")

                # instance of MIMEBase and named as p
                p = MIMEBase('application', 'octet-stream')

                # To change the payload into encoded form
                p.set_payload((attachment).read())

                # encode into base64
                encoders.encode_base64(p)

                p.add_header('Content-Disposition', "attachment; filename= %s" % filename)

                # attach the instance 'p' to instance 'msg'
                msg.attach(p)

        try:
            # creates SMTP session
            s = smtplib.SMTP('smtp.gmail.com', 587)

            # start TLS for security
            s.starttls()

            # Authentication
            s.login(fromaddr, password)

            # Converts the Multipart msg into a string
            text = msg.as_string()

            # sending the mail
            s.sendmail(fromaddr, receivers, text)

            # terminating the session
            s.quit()

            # IF THE MESSAGE WAS SUCCESSFULLY SENT
            return True
        except gaierror:
            pass
    
    def sendOutlook(self, receivers = [], subject = "Document for the last report", body = "Document for the last report", filenames = None):
        # Create an instance of the Outlook application
        outlook = win32.Dispatch('outlook.application')

        # Create a new email message
        mail = outlook.CreateItem(0)

        # Set the recipient, subject, and body of the email
        # nicole.bahati@africaimprovedfoods.com nsengumukiza_217236227@stud.ur.ac.rw
        mail.To = "; ".join(receivers) # 'nsengumukiza_217236227@stud.ur.ac.rw'
        mail.Subject = subject
        mail.Body = body

        # Add an attachment (optional)
        if filenames and receivers:
            for attachment in filenames:
                mail.Attachments.Add(attachment)

        # Send the email
        mail.Send()

        # IF THE MESSAGE WAS SUCCESSFULLY SENT
        return True
    
class PrintPaper:
    def __init__(self) -> None:
        pass

    def installed_printers(self):
        printers = win32print.EnumPrinters(win32print.PRINTER_ENUM_LOCAL | win32print.PRINTER_ENUM_CONNECTIONS, None, 1)

        printer_list = []
        for (flags, description, name, comment) in printers:
            printer_list.append(name)
        return printer_list

    def printPaper(self, printer, filename, extender = ['issue', 'troubleshout', 'machine'], kill_widget = None):
        """extender: files that are in path that I wished not to show the entire path, so I need to add their parent path here"""
        if filename[:filename.index('.')] in extender:
            filename = os.path.join(os.getcwd(), "docs", filename)
        
        # Get the default printer
        default_printer = win32print.GetDefaultPrinter()

        # Change the default printer to the specified printer
        win32print.SetDefaultPrinter(printer)
        win32api.ShellExecute (
        0,
        "print",
        filename,
        #
        # If this is None, the default printer will
        # be used anyway.
        #
        '/d:"%s"' % win32print.GetDefaultPrinter(),
        ".",
        0
        )

        # Set the default printer back to the original default printer
        win32print.SetDefaultPrinter(default_printer)

        if kill_widget:
            kill_widget.destroy()


def clipboard(data = None, action = None):
    if action == "copy" and data:
        copier.copy(data)
    elif action == "paste":
        return copier.paste()


def turn_into_secs(value, unit):
    """value: ex: 200
    unit: ex: days or hours and so on"""
    if 'sec' in unit:
        pass
    elif 'min' in unit:
        return 60 * value
    elif 'hour' in unit:
        return 3600 * value
    elif 'day' in unit:
        return 86400 * value
    elif 'year' in unit:
        return 86400 * 365 * value
    else:
        return None