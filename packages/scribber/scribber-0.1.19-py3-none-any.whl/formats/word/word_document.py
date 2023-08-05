from typing import Any

from pydantic import FilePath

from scribber.core.document import Title, Paragraph, EmptyLine, Table, Builder, CodeBlock
from docx import Document


class WordDocument:
    def __init__(self) -> None:
        self.parts = []
        self._report = Document()

    def add(self, part: Any) -> None:
        self.parts.append(part)

    def _render_title(self, title: Title) -> None:
        # str_len = len(title.title)
        self._report.add_heading(title.title, title.level - 1)

    def get_result(self) -> Document:
        for item in self.parts:
            if isinstance(item, Title):
                self._render_title(item)
            elif isinstance(item, Paragraph):
                self._render_paragraph(item)
            elif isinstance(item, EmptyLine):
                self._render_brake()
            elif isinstance(item, Table):
                self._render_table(item)
            elif isinstance(item, CodeBlock):
                self._render_code_block(item)
        return self._report

    def save(self, filename: FilePath):
        self.get_result().save(filename)

    def _render_paragraph(self, paragraph: Paragraph):
        self._report.add_paragraph(paragraph.text)

    def _render_code_block(self, code_block: CodeBlock):
        self._report.add_paragraph(code_block.code, style='Intense Quote')

    def _render_brake(self):
        self._report.add_paragraph("")

    def _render_table(self, item):
        table = self._report.add_table(rows=1, cols=len(item.headers))
        hdr_cells = table.rows[0].cells
        for i, hdr in enumerate(item.headers):
            hdr_cells[i].text = hdr
        for row in item.content:
            row_cells = table.add_row().cells
            for j, itm in enumerate(row):
                row_cells[j].text = str(itm)


class WordDocumentBuilder(Builder):
    def __init__(self) -> None:
        self.reset()

    def reset(self) -> None:
        self._word_report = WordDocument()

    @property
    def parts(self) -> WordDocument:
        parts = self._word_report
        self.reset()
        return parts

    def add_title(self, title: Title) -> None:
        self._word_report.add(title)

    def add_table(self, table: Table) -> None:
        self._word_report.add(table)

    def add_paragraph(self, paragraph: Paragraph) -> None:
        self._word_report.add(paragraph)

    def add_brake(self):
        self._word_report.add(EmptyLine())

    def add_code_block(self, code_block: CodeBlock):
        self._word_report.add(code_block)
